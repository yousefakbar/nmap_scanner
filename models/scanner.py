from PyQt5.QtWidgets import QMainWindow, QPushButton, QLineEdit, QGridLayout, QWidget, QLabel, QTextBrowser, \
        QScrollArea, QFrame, QMessageBox
from PyQt5.QtGui import QRegExpValidator, QDesktopServices, QIcon
from PyQt5.QtCore import QRegExp, Qt
from models.worker import Worker
from models.hosts import Host
from pymongo import MongoClient
import nmap
import re
import asyncio
import subprocess
import sys
import os
import platform
from datetime import datetime
from docx import Document

def ReverseString(text):
    return text[::-1]


class NmapScanner(QMainWindow):
    def __init__(self):
        super().__init__()
        self.layout = None

        self.api_key = '1f8da398-dc1d-4a00-be62-139a13cacda2'
        self.hosts_list = Host()
        self.host_widgets = []  # List to keep track of dynamically added labels
        self.layouts = {}  # Dictionary to keep track of QFrames and their layouts
        self.active_threads = []  # List to keep track of active threads
        self.buttons = [] # List to keep track of buttons
        self.result_objects = {}
        self.dynamic_buttons = []
        self.nvdlib_error = False
        self.nm_scan_error = False
        self.ports_scanned = False
        self.CVEs_checked = False
        self.import_data_files()
        self.init_ui()



    def import_data_files(self):
        if getattr(sys, 'frozen', False):
            application_path = sys._MEIPASS
            self.logo_ico = os.path.join(application_path, 'CoD_Logo.ico')
            self.logo_png = os.path.join(application_path, 'CoD_Logo.png')
        else:
            application_path = os.path.dirname(os.path.abspath(__file__))
            projectRoot = os.path.normpath(os.path.join(application_path, '..'))
            self.logo_ico = os.path.join(projectRoot, 'data/CoD_Logo.ico')
            self.logo_png = os.path.join(projectRoot, 'data/CoD_Logo.png')


    def init_ui(self):
        self.setIconFile()

        self.setWindowTitle('Nmap Scanner')
        self.setGeometry(100, 100, 600, 400)

        self.layout = QGridLayout()
        self.scroll_area = QScrollArea()
        self.scroll_area.setWidgetResizable(True)  # Allow the scroll area to resize the widget

        # Create a container widget to hold the layout
        container = QWidget()
        container.setLayout(self.layout)

        # Set the container widget as the widget for the scroll area
        self.scroll_area.setWidget(container)

        self.setCentralWidget(self.scroll_area)
        self.ip_input = QLineEdit(self)
        self.ip_input.setPlaceholderText('Enter IP or IP range')

        # IP Address Validation
        ip_regex = QRegExp("^(?:[0-9]{1,3}\.){3}[0-9]{1,3}$")
        ip_validator = QRegExpValidator(ip_regex, self.ip_input)
        self.ip_input.setValidator(ip_validator)

        self.layout.addWidget(self.ip_input, 0, 0, 1, 4)

        self.scan_button = QPushButton('Scan/فحص', self)
        self.scan_button.clicked.connect(lambda: self.on_scan_button_click(self.scan_button))
        self.layout.addWidget(self.scan_button, 1, 0, 1, 2)
        self.buttons.append(self.scan_button)


        # Add "Scan All In Network" button
        self.scan_all_buttons = QPushButton('Network Scan/فحص الشبكة', self)
        self.scan_all_buttons.clicked.connect(lambda: self.on_network_scan_button_click(self.scan_all_buttons))
        self.layout.addWidget(self.scan_all_buttons, 1, 2, 1, 2)
        self.buttons.append(self.scan_all_buttons)

        # Button to create reports
        self.create_report_button = QPushButton('Create Report', self)
        self.create_report_button.clicked.connect(self.create_report)
        self.layout.addWidget(self.create_report_button, 2, 0, 1, 2)
        self.buttons.append(self.create_report_button)
        self.create_report_button.setEnabled(False)  # disabled by default so no empty report is generated

        self.result_area = QTextBrowser(self)
        self.result_area.setReadOnly(True)
        self.result_area.setOpenExternalLinks(True)
        self.layout.addWidget(self.result_area, 3, 0, 1, 4)

        # "Reset" button
        self.reset_button = QPushButton('Reset/اعادة البرنامج', self)
        self.reset_button.clicked.connect(self.reset_view)
        self.layout.addWidget(self.reset_button, 2, 2, 1, 2)
        self.buttons.append(self.reset_button)

        # Check if nmap is installed. If not, ask user to install it first.
        self.__check_nmap_installation(container)


    def setIconFile(self):
        try:
            self.setWindowIcon(QIcon(self.logo_ico))
        except:
            print('There was an error setting the icon file')


    def __check_nmap_installation(self, container):
        if self.is_nmap_installed() == False:
            reply = QMessageBox.question(container, 'Install nmap?', 'Nmap is required to run the program. Would you like to install nmap?', QMessageBox.Yes | QMessageBox.No, QMessageBox.No)

            if reply == QMessageBox.No:
                QMessageBox.warning(container, 'nmap Required', 'nmap is required to run this command. Please install then run again.')
                sys.exit(1)
            else:
                self.install_nmap(container)


    async def perform_scan(self, ip=None):
        loop = asyncio.get_event_loop()
        if not ip:
            self.clear_dynamic_widgets()
            ip = self.ip_input.text()

        self.result_area.append('Scanning ports for IP: ' + ip)

        scanner = nmap.PortScanner()
        return await loop.run_in_executor(None, self.blocking_nmap_scan, scanner, ip, '-sC -sV')

    def display_results_perform_scan(self, scanner):
        self.clear_results()
        self.enable_all_buttons()
        self.create_report_button.setEnabled(True)

        if len(scanner.all_hosts()) == 0:
            self.result_area.append('IP was not found or is down. Please try another IP')
            return

        ip = scanner.all_hosts()[0]
        stats = scanner.scanstats()
        self.create_report_button.setEnabled(True)

        self.hosts_list.append_host(ip, scanner[ip].hostname())

        if self.nm_scan_error == True:
            return

        self.result_area.append('Port scan complete. Display open ports below.')

        numports = 0

        frame = QFrame(self)
        frame.setFrameShape(QFrame.StyledPanel)
        frame.setStyleSheet("background-color:gray")

        if getattr(sys, 'frozen', False):
            application_path = sys._MEIPASS
            sshFile = os.path.join(application_path, 'NMapScannerCSS.qss')
        else:
            application_path = os.path.dirname(os.path.abspath(__file__))
            application_path = os.path.normpath(os.path.join(application_path, '..'))
            sshFile = os.path.join(application_path, 'data/NMapScannerCSS.qss')
        with open(sshFile, "r") as fh:
            frame.setStyleSheet(fh.read())

        port_layout = QGridLayout(frame)

        sender = self.scanPressedButton  # Get the object that triggered the signal (in this case, the button)
        gridLayout = sender.parentWidget().layout()  # Get the grid layout

        row_position, column_position, _, _ = gridLayout.getItemPosition(gridLayout.indexOf(sender))

        if row_position == 1:
            row_position = 4
            column_position = 0
            self.scanPressedButton.setEnabled(True)
        else:
            self.layout.removeWidget(sender)
            self.buttons.remove(sender)
            self.dynamic_buttons.remove(sender)

            sender.deleteLater()  # Remove and delete the scan button

        self.layout.addWidget(frame, row_position, column_position, 1, 2)
        self.layouts[frame] = port_layout  # Store the QFrame and its layout in the dictionary
        for host in scanner.all_hosts():
            self.result_area.append(f'Host: {host} ({scanner[host].hostname()})')
            self.result_area.append(f'State: {scanner[host].state()}')

            for proto in scanner[host].all_protocols():
                lport = scanner[host][proto].keys()

                for port in lport:
                    numports = numports + 1

                    port_info = {
                            'product': scanner[ip]['tcp'][port]['product'],
                            'service': scanner[host][proto][port]["name"],
                            'version': scanner[host][proto][port]["version"]
                            }
                    self.hosts_list.append_port_to_host(host, port, port_info)

                    port_label = QLabel(self)
                    port_label.setText(
                        f'port: {port}\nstate:{scanner[host][proto][port]["state"]}\nservice: {scanner[host][proto][port]["name"]}\nversion: {scanner[host][proto][port]["version"]}')

                    version_scan_button = QPushButton(self)
                    version_scan_button.setText('Check for CVE')

                    port_layout.addWidget(port_label, numports, 0, 1, 1)
                    port_layout.addWidget(version_scan_button, numports, 1, 1, 1)

                    self.buttons.append(version_scan_button)
                    self.dynamic_buttons.append(version_scan_button)

                    version_scan_button.clicked.connect(
                        lambda checked, ip=host, inport=port: self.on_version_scan_button_click(ip, inport, version_scan_button))

            self.result_area.append('---------------------')

            # Delete this later
            print("TEST")
            self.hosts_list.print_dict()
            print("TEST2")


    async def perform_version_scan(self, ip, port):
        if self.CVEs_checked == False:
            self.CVEs_checked = True
        loop = asyncio.get_event_loop()
        args = '-sV -p ' + str(port)
        self.result_area.append("Scanning for CVE's on IP & port " + ip + ":" + str(port))
        scanner = nmap.PortScanner()
        return await loop.run_in_executor(None, self.blocking_version_scan, scanner, ip, port, args)

    def display_results_version_scan(self, scanner):
        self.enable_all_buttons()
        self.create_report_button.setEnabled(True)

        if self.nvdlib_error == True:
            print('There was an error in the nvdlib call')
            return

        if self.nm_scan_error == True:
            return

        self.result_area.append("Port version scan complete. CVE's shown below")

        ip = self.test_ip
        port = self.test_port
        sender = self.versionScanPressedButton  # Get the object that triggered the signal (in this case, the button)
        gridLayout = sender.parentWidget().layout()  # Get the grid layout

        row_position, _, _, _ = gridLayout.getItemPosition(gridLayout.indexOf(sender))

        self.layout.removeWidget(sender)
        self.buttons.remove(sender)
        self.dynamic_buttons.remove(sender)

        frame = next((key for key, value in self.layouts.items() if value == gridLayout), None)

        if frame:
            # Instantiate QTextBrowser with the correct parent (frame)
            self.portVulnTB = QTextBrowser(frame)
            self.portVulnTB.setStyleSheet("background-color: white; color: black")
            self.portVulnTB.setOpenExternalLinks(True)

            # Use the retrieved layout to add the QTextBrowser to the QFrame
            self.layouts[frame].addWidget(self.portVulnTB, row_position, 1, 1, 2)

        self.portVulnTB.append('Port: ' + str(port))
        self.portVulnTB.append('Service: ' + scanner[ip]['tcp'][port]['name'])
        self.portVulnTB.append('Name: ' + scanner[ip]['tcp'][port]['product'])
        self.portVulnTB.append('Version: ' + scanner[ip]['tcp'][port]['version'] + '\n')
        version = scanner[ip]['tcp'][port]['version']

        self.portVulnTB.append('Vulnerabilities:')
        self.portVulnTB.append('---------------\n')

        for line in self.version_res:
            self.portVulnTB.append(line)

        print("test3")
        self.hosts_list.print_dict()
        print("test4")

    def get_ssh_version(self, version):
        if len(version) <= 3:
            return version
        ssh_version = version[:version.index(' ')]
        match = re.search('[a-zA-Z]', ssh_version)
        match_idx = match.start()
        return ssh_version[:match_idx] + ' ' + ssh_version[match_idx:]

    def get_local_ip(self):
        os_type = platform.system()
        ip_cmd = {
                'Windows': 'ipconfig',
                'Linux': 'ip a',
                'Darwin': 'ifconfig en0' # MacOS -- Works only for wireless interface
                }.get(os_type)

        if not ip_cmd:
            print('Unsupported Operating System')
            raise ValueError('Unsupported Operating System')

        result = subprocess.run(ip_cmd, stdout=subprocess.PIPE, text=True, shell=True)
        output = result.stdout

        # For private IP addresses, these will typically be the range in which
        # the devices are on in the network. So, look for patterns in the ip
        # program that match the following Private IP Address Ranges:
        # - 10.0.0.0 to 10.255.255.255
        # - 172.16.0.0 to 172.31.255.255
        # - 192.168.0.0 to 192.168.255.255
        ip_patterns = {
                'Windows': r'IPv4 Address[^:]*:\s*(\d+\.\d+\.\d+\.\d+)',
                'Linux': r'inet (\d+\.\d+\.\d+\.\d+)/\d+',
                'Darwin': r'inet (\d+\.\d+\.\d+\.\d+) '  # Same as Linux, but keeping separate for clarity
                }

        pattern = ip_patterns[os_type]
        ips = re.findall(pattern, output)

        # filter out the loopback IP and get the private IP in the network
        private_ip = [ip for ip in ips if not ip.startswith('127.') and
                (ip.startswith('10.') or ip.startswith('172.') or
                    ip.startswith('192.168.'))]

        return private_ip[0]

    async def scan_all_in_network(self):
        self.clear_dynamic_widgets()
        loop = asyncio.get_event_loop()

        if self.ip_input.text() == '':
            ip = self.get_local_ip()
        else:
            ip = self.ip_input.text()

        if ip:  # Use the provided IP to determine the network range
            network = '.'.join(ip.split('.')[:3]) + '.0/24'  # Assumes a class C network
            self.result_area.append('Scanning network IP: ' + network)
            scanner = nmap.PortScanner()
            return await loop.run_in_executor(None, self.blocking_nmap_scan, scanner, network, '-sn')
        else:
            self.result_area.setText("Please enter a valid IP address to determine the network.")

    def display_results_scan_all(self, scanner):
        self.enable_all_buttons()
        self.result_area.clear()
        self.scanAllPressedButton.setEnabled(True)
        self.create_report_button.setEnabled(True)
        i = 3  # Adjust if more fields are added above the results field
        scan_stats = scanner.scanstats()
        self.result_area.append('Network scan complete. Displaying hosts below.\n')
        self.result_area.append('List of hosts UP (%s/%s) in network \n' % (scan_stats['uphosts'], scan_stats['totalhosts']))
        for host in scanner.all_hosts():  # for each host found, create a qLabel and "more" button
            i += i
            self.hosts_list.append_host(host, scanner[host].hostname())

            self.hostLabel = QLabel(self)
            self.hostLabel.setText(f'IP: {host}\nHostname: {scanner[host].hostname()}\n')
            print(scanner[host])
            self.hostLabel.setAlignment(Qt.AlignTop)

            self.performScanButton = QPushButton(self)
            self.performScanButton.setText('Scan Ports')

            self.buttons.append(self.performScanButton)
            self.dynamic_buttons.append(self.performScanButton)

            self.layout.addWidget(self.hostLabel, i, 0, 1, 2)
            self.layout.addWidget(self.performScanButton, i, 2, 1, 1)
            self.host_widgets.append(self.hostLabel)


            self.performScanButton.clicked.connect(lambda checked, ip=host: self.on_scan_button_click(self.performScanButton, ip))


    def blocking_nmap_scan(self, nm, ip, args):
        try:
            nm.scan(ip, arguments=args)
        except:
            print('There was an error in the nmap scan. Try again.')
            self.nm_scan_error = True
        return nm

    def blocking_version_scan(self, scanner, ip, port, args):
        try:
            scanner.scan(ip, arguments=args)
        except:
            print('There was an error in the nmap version scan. Try again.')
            self.nm_scan_error = True
            return

        self.version_res = []

        service = scanner[ip]['tcp'][port]['product']
        if scanner[ip]['tcp'][port]['product'] == 'OpenSSH':
            version = self.get_ssh_version(scanner[ip]['tcp'][port]['version'])
        else:
            version = scanner[ip]['tcp'][port]['version']

        try:
            client = MongoClient() # TODO modify based on server DB credentials
            db = client.cvedb
            cpes = db.cpe
            query_cpe = { 'product': service.lower(), 'version': version}
            results = cpes.find(query_cpe)
            cpe_list = []
            for cpe in results:
                cpe_list.append(cpe)
            cpe_name = cpe_list[0]['cpeName']
            print('cpe name: ' + cpe_name)
            print('service: ' + service.lower())
            print('version: ' + version)
        except:
            print('error searching cpe')
            self.nvdlib_error = True
            return scanner

        try:
            client = MongoClient()
            db = client.cvedb
            cves = db.cves
            query_cve = { 'vulnerable_configuration': cpe_name }
            results = cves.find(query_cve)
            for cve in results:
                self.version_res.append('Severity: ' + str(cve['cvss3']))
                self.version_res.append('<a href="https://nvd.nist.gov/vuln/detail/' + str(cve['id']) + '">' + str(cve['id']) + '</a>')
                self.version_res.append('Last Updated: ' + str(cve['lastModified']))
                self.version_res.append('')
                cve_info = {
                        'last_modified': cve['lastModified'],
                        'url': 'https://nvd.nist.gov/vuln/detail/' + cve['id'],
                        'severity': cve['cvss3'],
                        }
                self.hosts_list.append_cve_to_port(ip, port, cve['id'], cve_info)
        except:
            self.nvdlib_error = True
            return scanner

        self.test_ip = ip
        self.test_port = port
        return scanner

    def clear_results(self):
        self.result_area.clear()
        self.result_objects.clear()

    def clear_dynamic_widgets(self):
        # Remove and delete each dynamically added widget

        for label in self.host_widgets:


            self.layout.removeWidget(label)
            label.deleteLater()
        self.host_widgets.clear()  # Clear the list after removing the labels

        for button in self.dynamic_buttons:
            self.layout.removeWidget(button)
            button.deleteLater()
        self.dynamic_buttons.clear() # clear it after removing all buttons

        # Remove and Delete each frame and it's layout
        for frame, layout in self.layouts.items():
            self.layout.removeWidget(frame)
            frame.deleteLater()

            # Remove and delete widgets from the layout
            for i in reversed(range(layout.count())):
                layout_item = layout.itemAt(i)
                if layout_item is not None:
                    widget = layout_item.widget()
                    if widget is not None:
                        layout.removeWidget(widget)
                        widget.deleteLater()

        self.result_objects.clear()
        self.layouts.clear()  # Clear the dictionary after removing the frames and layouts

    def reset_view(self):
        self.clear_dynamic_widgets()
        self.clear_results()
        self.ip_input.setText('')
        self.ports_scanned = False
        self.CVEs_checked = False
        self.create_report_button.setEnabled(False)  # disable the button so you cannot generate a blank report

    def on_scan_button_click(self, button, ip=None):
        self.disable_all_buttons()
        self.scanPressedButton = self.sender()
        if self.ports_scanned == False:
            self.ports_scanned = True

        if ip: # This will run when the user clicks scan host after "Scanning all in network"
            self.start_async_task(self.perform_scan(ip), self.display_results_perform_scan)
        elif button == self.scan_button and self.ip_input.text(): # This will run when the user enters IP in ipInput and clicks Scan
            self.start_async_task(self.perform_scan(), self.display_results_perform_scan)
        else: # This will run when the user clicks the scan button with empty input field
            self.result_area.append('Please enter a valid IP to scan')
            self.enable_all_buttons()

    def on_network_scan_button_click(self, button):
        self.disable_all_buttons()
        self.scanAllPressedButton = self.sender()
        self.scanAllPressedButton.setEnabled(False)
        self.start_async_task(self.scan_all_in_network(), self.display_results_scan_all)

    def on_version_scan_button_click(self, ip, port, button):
        for b in self.dynamic_buttons:
            if b != self.sender() and isinstance(b, QPushButton) and b.text() == 'Check for CVE':
                pass
            else:
                b.setEnabled(False)

        self.versionScanPressedButton = self.sender()
        self.versionScanPressedButton.setEnabled(False)
        self.start_async_task(self.perform_version_scan(ip, port), self.display_results_version_scan) # TODO

    def start_async_task(self, coroutine, callback):
        task = Worker(coroutine)
        task.scan_complete.connect(callback)
        task.finished.connect(lambda: self.active_threads.remove(task))  # Connect to a slot to remove the thread from the list
        self.active_threads.append(task)  # Add the thread to the list of active threa
        task.start()

    def is_nmap_installed(self):
        return subprocess.run(['which', 'nmap'], stdout=subprocess.PIPE).returncode == 0

    def install_nmap(self, container):
        QMessageBox.information(container, 'Installing', 'Please wait while we download and install nmap')
        try:
            subprocess.run(['pkexec', 'sudo', 'apt', 'install', '-y', 'nmap'], check=True)
            QMessageBox.information(container, 'nmap Installed', 'nmap has been successfully installed. You may now use the program')
        except:
            QMessageBox.critical(container, 'Installation failed', 'Installation failed.')

    # functions to disable and enable buttons
    def disable_all_buttons(self):
        for i in range(4):
            self.buttons[i].setEnabled(False)
            i = i + 1
        for button in self.dynamic_buttons:
            if button:
                button.setEnabled(False)

    def enable_all_buttons(self):

        for i in range(4):
            self.buttons[i].setEnabled(True)
            i=i+1
        for button in self.dynamic_buttons:
            button.setEnabled(True)

        if len(self.result_objects) == 0:
            self.create_report_button.setEnabled(False) # if there are no results, disable report creation

    def create_report(self):

        print('Generating report')
        document = Document('Single_Scan_Template.docx') # opens template

        # Generates title for document
        now = datetime.now()
        ip_address = str(list(self.hosts_list.hosts_list.keys())[0])
        title = ip_address + '_' + now.strftime("%d.%m.%Y_%H.%M.%S") + '.docx'

        document.paragraphs[0].add_run(ip_address) # adds to title


        # t1 is the table of ips and host names
        num_hosts = len(self.hosts_list.hosts_list)+1
        t1 = document.add_table(rows=num_hosts, cols=2)


        for x in range(num_hosts):
            if x == 0: # fills first row
                t1.rows[x].cells[0].text = "IP Address"
                t1.rows[x].cells[1].text = "Host name"
            else: # fills rest of the table
                ip_address = str(list(self.hosts_list.hosts_list.keys())[x-1])
                ip_hostname = self.hosts_list.hosts_list[ip_address].hostname
                t1.rows[x].cells[0].text = ip_address
                t1.rows[x].cells[1].text = ip_hostname

        if self.ports_scanned:
            print('ports scanned')
            document.add_paragraph().text = "\nScanned ports are below: \n"
            # for each IP, make a table for its ports

            for x in range(num_hosts-1):
                ip_address = str(list(self.hosts_list.hosts_list.keys())[x])
                num_ports = len(self.hosts_list.hosts_list[ip_address].ports)
                print(num_ports)
                if num_ports:
                    t2 = document.add_table(rows=num_ports+1,cols=6)
                    for i in range(num_ports+1):
                        if i == 0: # Fills first row in table
                            t2.rows[i].cells[0].text = "IP Address"
                            t2.rows[i].cells[1].text = "Host name"
                            t2.rows[i].cells[2].text = "Port number"
                            t2.rows[i].cells[3].text = "Product"
                            t2.rows[i].cells[4].text = "Service name"
                            t2.rows[i].cells[5].text = "Service version"
                        else: # fills rest of table
                            ip_address = str(list(self.hosts_list.hosts_list.keys())[x])
                            temp_ip = self.hosts_list.hosts_list[ip_address]
                            ip_hostname = temp_ip.hostname
                            list_of_ports = list(temp_ip.ports.keys())
                            t2.rows[i].cells[0].text = ip_address
                            t2.rows[i].cells[1].text = ip_hostname
                            t2.rows[i].cells[2].text = str(temp_ip.ports[list_of_ports[i-1]].port_number)
                            t2.rows[i].cells[3].text = temp_ip.ports[list_of_ports[i-1]].product
                            t2.rows[i].cells[4].text = temp_ip.ports[list_of_ports[i-1]].service_name
                            t2.rows[i].cells[5].text = temp_ip.ports[list_of_ports[i-1]].service_version

                else:
                    document.add_paragraph().text = "No ports found in " + ip_address

        if self.CVEs_checked:
            for x in range(num_hosts-1):
                ip_address = str(list(self.hosts_list.hosts_list.keys())[x])
                num_ports = len(self.hosts_list.hosts_list[ip_address].ports)
                temp_ip = self.hosts_list.hosts_list[ip_address]
                list_of_ports = list(temp_ip.ports.values())
                for p in range(num_ports):
                    temp_port = temp_ip.ports[list_of_ports[p].port_number]
                    document.add_paragraph('Port: ' + str(temp_port) + ":\n")
                    num_cves = len(temp_ip.ports[list_of_ports[p]].port.cves)

                    t3 = document.add_table(rows=num_cves+1,cols=4)
                    for i in range(num_cves+1):
                        if i == 0:
                            t3.rows[i].cells[0].text = 'CVE'
                            t3.rows[i].cells[1].text = 'Last modified'
                            t3.rows[i].cells[2].text = 'URL'
                            t3.rows[i].cells[3].text = 'Severity'
                        else:
                            list_of_cves = list(temp_port.cves.keys())
                            t3.rows[i].cells[0].text = temp_port.port.cves[list_of_cves[i-1]].name
                            t3.rows[i].cells[1].text = temp_port.port.cves[list_of_cves[i-1]].last_modified
                            t3.rows[i].cells[2].text = temp_port.port.cves[list_of_cves[i - 1]].url
                            t3.rows[i].cells[3].text = temp_port.port.cves[list_of_cves[i - 1]].severity


        document.save(title)


